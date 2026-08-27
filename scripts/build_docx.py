#!/usr/bin/env python3
"""Render a tailored resume.md into resume.docx, matching the PDF layout.

Companion to build_resume.py. Same markdown in, same visual design out -- Georgia
10pt, 0.6/0.7in margins, centered header, rust section headings with a hairline
rule, right-aligned locations and dates. The difference is that a .docx is
editable, which is what recruiters, staffing agencies and career coaches ask for
when they want to add their own header or tweak wording.

Parsing is imported from build_resume.py rather than reimplemented. That file
carries hard-won rules about which lines are job titles and which are bullets,
and two resumes once shipped with no employment dates because a parser branch was
too narrow. One parser, one set of rules.

Usage:
    python3 scripts/build_docx.py user/tailored/misc/randstad-career-services
    python3 scripts/build_docx.py --all
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_resume import ROLE_LINE, parse  # noqa: E402

ROOT = Path(__file__).resolve().parent.parent

FONT = "Georgia"
BODY_PT = 10
ACCENT = RGBColor(0xB4, 0x4A, 0x1E)
CONTENT_WIDTH = Inches(7.1)  # letter 8.5in less 0.7in margins each side


def load_config():
    cfg = {}
    for line in (ROOT / "user" / "config.yaml").read_text().split("\n"):
        m = re.match(r'^(\w+):\s*"?([^"#]*?)"?\s*(?:#.*)?$', line)
        if m and m.group(2):
            cfg[m.group(1)] = m.group(2).strip()
    return cfg


def add_border(paragraph, edge="bottom", size=6, color="CCCCCC"):
    """Hairline rule under a paragraph. python-docx has no API for this."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), "1")
    el.set(qn("w:color"), color)
    borders.append(el)
    pPr.append(borders)


def write_runs(paragraph, text, *, size=BODY_PT, italic=False):
    """Emit markdown inline formatting as runs. Bold and links; links keep their text.

    DO NOT add non-breaking hyphens here. build_resume.py wraps hyphenated tokens in
    white-space:nowrap because a hyphen that falls at a PDF line break extracts without
    the hyphen, so an ATS reads "endtoend". The docx has the same rendering behaviour --
    verified 2026-08-22, LibreOffice rendered this resume and pdftotext lost the hyphen
    from "end-to-end" and "re-pushed".

    It is still the wrong fix here. Measured both alternatives: U+2011 and OOXML
    w:noBreakHyphen each stop the break, and each extracts as U+2011 rather than an
    ASCII hyphen, so a search for "end-to-end" fails either way -- one broken match
    traded for another. And a .docx is normally parsed as OOXML, not rendered first:
    all 30 hyphenated tokens in this resume are contiguous and ASCII in
    word/document.xml, where line breaking does not exist. Substituting U+2011 would
    corrupt that primary path to protect a secondary one.
    """
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = i % 2 == 1
        run.italic = italic
    return paragraph


def para(doc, *, space_before=0, space_after=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    return p


def two_column(doc, left, right, *, left_bold=False, left_italic=False,
               right_italic=False, left_size=BODY_PT, space_before=0):
    """Left text with a right-aligned tab stop at the margin, as in the PDF."""
    p = para(doc, space_before=space_before)
    p.paragraph_format.tab_stops.add_tab_stop(CONTENT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(left)
    run.font.name, run.font.size = FONT, Pt(left_size)
    run.bold, run.italic = left_bold, left_italic
    run = p.add_run("\t" + right)
    run.font.name, run.font.size = FONT, Pt(BODY_PT)
    run.italic = right_italic
    return p


def section_heading(doc, name):
    p = para(doc, space_before=8, space_after=2)
    run = p.add_run(name.upper())
    run.font.name, run.font.size = FONT, Pt(11)
    run.bold = True
    run.font.color.rgb = ACCENT
    add_border(p)
    return p


def bullet(doc, text):
    p = para(doc, space_after=2)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.13)
    run = p.add_run("•  ")
    run.font.name, run.font.size = FONT, Pt(BODY_PT)
    write_runs(p, text)
    return p


def render_experience(doc, lines):
    dropped = []
    for line in lines:
        s = line.strip()
        if not s or s == "---":
            continue
        if s.startswith("### "):
            company, _, loc = s[4:].partition("—")
            two_column(doc, company.strip(), loc.strip(), left_bold=True,
                       right_italic=True, left_size=10.5, space_before=6)
        elif s.startswith("- "):
            bullet(doc, s[2:])
        elif ROLE_LINE.search(s):
            title, _, dates = s.partition("—")
            two_column(doc, re.sub(r"\*+", "", title).strip(), dates.strip(),
                       left_italic=True)
        elif s.startswith("*"):
            raise SystemExit(
                f"  subsection labels are not used -- delete this line: {s!r}\n"
                f"  see user/resume-style.md, 'No subsection labels inside Experience'")
        else:
            dropped.append(s)
    if dropped:
        print("WARNING: lines in ## Experience matched no rule and were NOT rendered:",
              file=sys.stderr)
        for s in dropped:
            print(f"    {s[:100]}", file=sys.stderr)


def build(target: Path):
    md = (target / "resume.md").read_text()
    cfg = load_config()

    title_line = re.search(r"^\*\*(.+?)\*\*\s*(?:\||$)", md, re.M)
    if not title_line:
        print("ERROR: no headline found. Expected a line like "
              "'**Senior Software Engineer**' near the top of resume.md.", file=sys.stderr)
        raise SystemExit(1)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = FONT, Pt(BODY_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.7)

    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    run = p.add_run(cfg["name"])
    run.font.name, run.font.size, run.bold = FONT, Pt(22), True

    contact = " · ".join([title_line.group(1).strip(), cfg["location"],
                               cfg["email"], cfg["phone"]])
    for text in (contact, f'{cfg["linkedin"]} | {cfg["github"]}'):
        write_runs(para(doc, align=WD_ALIGN_PARAGRAPH.CENTER), text)

    add_border(para(doc, space_before=4, space_after=4), color="000000", size=8)

    for name, lines in parse(md):
        body = [l for l in lines if l.strip() and l.strip() != "---"]
        section_heading(doc, name)
        if name.lower() == "experience":
            render_experience(doc, lines)
        elif name.lower() == "skills":
            for l in body:
                if l.strip().startswith("- "):
                    write_runs(para(doc, space_after=1), l.strip()[2:])
        elif any(l.strip().startswith("- ") for l in body):
            for l in body:
                if l.strip().startswith("- "):
                    bullet(doc, l.strip()[2:])
        else:
            write_runs(para(doc, space_after=2), " ".join(body))

    out = target / "resume.docx"
    doc.save(str(out))
    return out


def main():
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    if "--all" in sys.argv:
        targets = sorted(p.parent for p in (ROOT / "user" / "tailored").rglob("resume.md"))
    elif args:
        targets = [Path(args[0])]
    else:
        print(__doc__)
        return 1

    for target in targets:
        if not (target / "resume.md").exists():
            print(f"ERROR: no resume.md in {target}", file=sys.stderr)
            return 1
        out = build(target)
        print(f"{target.name}: wrote {out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
